# -*- coding: utf-8 -*-
"""Build b11402037-b11402053-report.docx — C Programming term project report."""
from docx import Document
from docx.shared import Pt, Mm, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ZH = "Microsoft JhengHei"
MONO = "Consolas"
NAVY = RGBColor(0x1B, 0x2A, 0x41)
TEAL = RGBColor(0x0E, 0x8C, 0x8B)
INK = RGBColor(0x22, 0x30, 0x3F)
MUTED = RGBColor(0x5F, 0x6E, 0x7D)
CODE_FILL = "F2F5F7"
HDR_FILL = "1B2A41"

doc = Document()

# ---------- page setup: A4, 2.54cm margins ----------
sec = doc.sections[0]
sec.page_width = Mm(210)
sec.page_height = Mm(297)
sec.top_margin = sec.bottom_margin = Cm(2.54)
sec.left_margin = sec.right_margin = Cm(2.54)
sec.different_first_page_header_footer = True

# ---------- base styles ----------
def style_font(style, size, bold=False, color=None, font=ZH):
    style.font.name = font
    style.font.size = Pt(size)
    style.font.bold = bold
    if color is not None:
        style.font.color.rgb = color
    rPr = style.element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts'); rPr.append(rFonts)
    rFonts.set(qn('w:eastAsia'), ZH)

style_font(doc.styles['Normal'], 11, color=INK)
doc.styles['Normal'].paragraph_format.line_spacing = 1.4
doc.styles['Normal'].paragraph_format.space_after = Pt(6)

style_font(doc.styles['Heading 1'], 16, bold=True, color=NAVY)
doc.styles['Heading 1'].paragraph_format.space_before = Pt(18)
doc.styles['Heading 1'].paragraph_format.space_after = Pt(8)
style_font(doc.styles['Heading 2'], 13, bold=True, color=TEAL)
doc.styles['Heading 2'].paragraph_format.space_before = Pt(12)
doc.styles['Heading 2'].paragraph_format.space_after = Pt(6)

# ---------- helpers ----------
def set_run(run, size=11, bold=False, italic=False, color=INK, font=ZH):
    run.font.name = font
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = color
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts'); rPr.append(rFonts)
    rFonts.set(qn('w:eastAsia'), ZH)

def para(text=None, style=None, align=None, runs=None, space_after=None):
    p = doc.add_paragraph(style=style)
    if align is not None:
        p.alignment = align
    if space_after is not None:
        p.paragraph_format.space_after = Pt(space_after)
    if runs:
        for txt, opts in runs:
            set_run(p.add_run(txt), **opts)
    elif text is not None:
        set_run(p.add_run(text))
    return p

def shade_para(p, fill=CODE_FILL):
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear'); shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill)
    pPr.append(shd)

def code_block(lines, size=9.5):
    for i, ln in enumerate(lines):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.15
        p.paragraph_format.left_indent = Cm(0.3)
        set_run(p.add_run(ln if ln else " "), size=size, font=MONO, color=INK)
        shade_para(p)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)

def shade_cell(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear'); shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill)
    tcPr.append(shd)

def make_table(headers, rows, widths_cm, mono_cols=()):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = 'Table Grid'
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.autofit = False
    for j, w in enumerate(widths_cm):
        for r in range(1 + len(rows)):
            t.cell(r, j).width = Cm(w)
    for j, h in enumerate(headers):
        c = t.cell(0, j)
        shade_cell(c, HDR_FILL)
        p = c.paragraphs[0]; p.paragraph_format.space_after = Pt(0)
        set_run(p.add_run(h), size=10, bold=True, color=RGBColor(0xFF, 0xFF, 0xFF))
    for i, row in enumerate(rows):
        for j, val in enumerate(row):
            c = t.cell(i + 1, j)
            if i % 2 == 1:
                shade_cell(c, "F0F5F7")
            p = c.paragraphs[0]; p.paragraph_format.space_after = Pt(0)
            set_run(p.add_run(val), size=9.5,
                    font=MONO if j in mono_cols else ZH)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return t

def figure(path, caption, width_cm=15.9):
    doc.add_picture(path, width=Cm(width_cm))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.paragraphs[-1].paragraph_format.space_before = Pt(6)
    para(align=WD_ALIGN_PARAGRAPH.CENTER,
         runs=[(caption, dict(size=10, color=MUTED))], space_after=10)

def bullets(items):
    for it in items:
        p = doc.add_paragraph(style='List Bullet')
        p.paragraph_format.space_after = Pt(3)
        if isinstance(it, list):
            for txt, opts in it:
                set_run(p.add_run(txt), **opts)
        else:
            set_run(p.add_run(it))

# ---------- footer page number (non-first pages) ----------
footer = sec.footer
fp = footer.paragraphs[0]
fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
r1 = fp.add_run(); set_run(r1, size=9, color=MUTED)
fld1 = OxmlElement('w:fldChar'); fld1.set(qn('w:fldCharType'), 'begin')
r1._element.append(fld1)
r2 = fp.add_run(); set_run(r2, size=9, color=MUTED)
instr = OxmlElement('w:instrText'); instr.set(qn('xml:space'), 'preserve')
instr.text = ' PAGE '
r2._element.append(instr)
r3 = fp.add_run(); set_run(r3, size=9, color=MUTED)
fld2 = OxmlElement('w:fldChar'); fld2.set(qn('w:fldCharType'), 'end')
r3._element.append(fld2)

# =====================================================================
# Cover
# =====================================================================
for _ in range(4):
    doc.add_paragraph()
para("C Programming / C 程式設計", align=WD_ALIGN_PARAGRAPH.CENTER,
     runs=[("C Programming / C 程式設計", dict(size=14, color=MUTED))])
para(align=WD_ALIGN_PARAGRAPH.CENTER,
     runs=[("期末專題報告 Term Project Report", dict(size=14, color=MUTED))])
doc.add_paragraph()
para(align=WD_ALIGN_PARAGRAPH.CENTER,
     runs=[("Clawer C Data Normalization Engine", dict(size=24, bold=True, color=NAVY))])
para(align=WD_ALIGN_PARAGRAPH.CENTER,
     runs=[("C 語言大學排名資料正規化引擎", dict(size=18, bold=True, color=TEAL))])
for _ in range(3):
    doc.add_paragraph()

# member table on cover
t = doc.add_table(rows=3, cols=2)
t.style = 'Table Grid'
t.alignment = WD_TABLE_ALIGNMENT.CENTER
cover_rows = [("學號", "姓名"), ("b11402037", "高恩在"), ("b11402053", "林秉學")]
for r in range(3):
    for c_i in range(2):
        cell = t.cell(r, c_i)
        cell.width = Cm(4.5)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        if r == 0:
            shade_cell(cell, HDR_FILL)
            set_run(p.add_run(cover_rows[r][c_i]), size=11, bold=True,
                    color=RGBColor(0xFF, 0xFF, 0xFF))
        else:
            set_run(p.add_run(cover_rows[r][c_i]), size=12,
                    font=MONO if c_i == 0 else ZH)
for _ in range(3):
    doc.add_paragraph()
para(align=WD_ALIGN_PARAGRAPH.CENTER,
     runs=[("2026 年 6 月", dict(size=12, color=MUTED))])

doc.add_page_break()

# =====================================================================
# 1. Introduction
# =====================================================================
doc.add_heading("1. Introduction｜簡介", level=1)

doc.add_heading("1.1 背景與動機", level=2)
para("上游爬蟲系統自 QS World University Rankings、Times Higher Education（THE）與 "
     "Shanghai Ranking（ARWU）等多個大學排名來源抓取資料。由於各來源的欄位慣例不一致，"
     "爬蟲輸出的 CSV 格式高度異質，無法直接進行跨來源的比較與統計分析。主要的異質性來源包括：")
bullets([
    [("大學名稱：", dict(bold=True)), ("大小寫不一、知名縮寫（MIT、ETH、KAUST 等）、標點與多餘空白差異。", dict())],
    [("國家名稱：", dict(bold=True)), ("縮寫（U.S.A.、UK、PRC）、官方長名（Federal Republic of Germany）與「The X」前綴變體（The Netherlands）。", dict())],
    [("排名格式：", dict(bold=True)), ("純數字（53）、區間（101-150）、文字前綴（Top 100、Rank 53、=201、#10、201+）等多種寫法。", dict())],
    [("分數欄位：", dict(bold=True)), ("含前後空白、前導小數點（.5）、空字串與 N/A 等無效值。", dict())],
])

doc.add_heading("1.2 研究問題", level=2)
p = para(runs=[("在不修改下游分析邏輯的前提下，如何將異質性爬蟲輸出（多種大學名稱格式、"
                "排名區間字串、不一致的國家縮寫）轉換為可直接比較的數值紀錄？",
                dict(bold=True, color=NAVY))])
shade_para(p, "E8F4F2")

doc.add_heading("1.3 設計選擇", level=2)
para("本專題以 C 語言（C11，gcc -Wall -Wextra）實作核心引擎，而非採用 Python / pandas。"
     "主要考量為：大量 CSV 批次處理時可維持低記憶體佔用與可預期的執行時間；"
     "且無 Python 執行期依賴，便於與上游爬蟲管線的部署整合。")

doc.add_heading("1.4 系統範圍", level=2)
bullets([
    "輸入：爬蟲輸出 CSV（11 欄原始格式）。",
    "輸出：標準化 CSV（5 欄：University, Country, Rank Min, Rank Max, Overall Score）。",
    "邊界：單機、檔案導向；無資料庫伺服器、無 Web API、無外部網路服務依賴。",
])

# =====================================================================
# 2. System Design
# =====================================================================
doc.add_heading("2. System Design｜系統設計", level=1)

doc.add_heading("2.1 系統架構與資料流", level=2)
para("系統採單向資料流的 pipeline 架構：CSV Reader 讀入原始資料並填入記憶體中的 "
     "UniversityRecord 陣列，Normalization Pipeline 依序呼叫四個解析模組就地正規化，"
     "最後由 CSV Writer 輸出標準化結果。各模組以標頭檔（include/）定義公開介面，彼此低耦合：")
code_block([
    "raw_universities.csv",
    "    → csv_reader.c      （header 驗證 / quoted CSV 解析 / BOM 清理）",
    "    → UniversityRecord[]（raw 欄位）",
    "    → normalizer.c      （pipeline 協調器）",
    "         ├─ name_normalizer.c     名稱清理與縮寫保留",
    "         ├─ country_normalizer.c  國家別名映射",
    "         ├─ rank_parser.c         排名區間解析",
    "         └─ score_parser.c        分數提取",
    "    → csv_writer.c      （quoting / escaping）",
    "    → normalized_universities.csv",
])
make_table(
    ["模組", "檔案", "主要責任"],
    [
        ("CLI 主程式", "src/main.c", "互動式選單，協調載入、正規化、預覽、匯出"),
        ("CSV Reader", "src/csv_reader.c", "header 驗證、quoted CSV 解析、BOM 清理、不合法列略過"),
        ("Record 模型", "include/record.h", "保存 raw 與 normalized 欄位、rank_min/max、score"),
        ("Pipeline 協調器", "src/normalizer.c", "逐筆呼叫名稱、國家、排名、分數四個解析模組"),
        ("Name Normalizer", "src/name_normalizer.c", "名稱清理、大小寫調整、知名縮寫保留"),
        ("Country Normalizer", "src/country_normalizer.c", "國家別名映射與標準化（38+ 條目）"),
        ("Rank Parser", "src/rank_parser.c", "解析 Top 100、Rank 53、101-150、=201、#10、201+"),
        ("Score Parser", "src/score_parser.c", "字串分數提取，失敗回傳哨兵值 -1"),
        ("CSV Writer", "src/csv_writer.c", "標準化輸出，必要時 quoting / escaping"),
        ("共用工具", "src/utils.c", "trim、collapse spaces、title case、縮寫偵測"),
    ],
    [2.8, 5.4, 7.8], mono_cols=(1,))

doc.add_heading("2.2 CSV Schema 設計", level=2)
para("本系統採檔案導向設計，以 CSV 作為持久層，不依賴外部資料庫伺服器，"
     "降低部署複雜度並貼合爬蟲批次輸出的使用情境。輸入與輸出格式如下：")
code_block([
    "輸入（11 欄）：QS Rank,University,Country,GMAT,GRE,GPA,IELTS,TOEFL,",
    "              Duolingo,Overall Score,URL",
    "輸出（ 5 欄）：University,Country,Rank Min,Rank Max,Overall Score",
])
para("實際資料範例（正規化前 → 後）：")
code_block([
    'raw  | University of   California Berkeley (UCB)  , usa ,12 , 14 ,  91.7',
    'norm | University of California Berkeley (Ucb),United States,12,14,91.70',
])

doc.add_heading("2.3 核心資料結構", level=2)
para("UniversityRecord 同時保存 raw 與 normalized 欄位，便於前後對照、預覽與除錯：")
code_block([
    "typedef struct {",
    "    /* Raw input fields */",
    "    char raw_name[150];",
    "    char raw_country[80];",
    "    char raw_rank[50];",
    "    char raw_score[50];",
    "    /* Normalized fields */",
    "    char normalized_name[150];",
    "    char normalized_country[80];",
    "    int    rank_min, rank_max;",
    "    double score;",
    "} UniversityRecord;",
])

doc.add_heading("2.4 關鍵設計決策", level=2)
bullets([
    [("國家別名對照表：", dict(bold=True)),
     ("靜態 key → value 查表。先以 normalize_basic（去標點、轉小寫、壓縮空白）做正規化，"
      "再比對別名表；涵蓋縮寫、官方長名與「The X」前綴變體，目前共 38+ 條目。", dict())],
    [("排名區間表示：", dict(bold=True)),
     ("以 rank_min / rank_max 兩個整數欄位統一表達單值與區間"
      "（53 → [53,53]；101-150 → [101,150]；Top 100 → [1,100]），下游可直接做數值比較。", dict())],
    [("哨兵值設計：", dict(bold=True)),
     ("無法解析的分數標記為 -1，由 writer 層轉為空欄位輸出，避免髒資料混入下游統計。", dict())],
    [("Reader 防衛機制：", dict(bold=True)),
     ("header 欄位數與欄名驗證（不符即拒絕載入）、RFC 4180 quoted CSV 解析、"
      "UTF-8 BOM 清理、不合法列略過。", dict())],
])

# =====================================================================
# 3. Implementation
# =====================================================================
doc.add_heading("3. Implementation｜實作", level=1)

doc.add_heading("3.1 開發環境與建置", level=2)
para("開發語言為 C11，編譯器為 gcc（編譯選項 -Wall -Wextra -std=c11），"
     "以 Makefile 管理主程式與九個測試套件的建置。"
     "macOS / Linux 環境直接使用 make；Windows 環境（MSYS2 gcc）可用以下指令直接編譯：")
code_block([
    "gcc -Wall -Wextra -std=c11 -Iinclude src/main.c src/normalizer.c \\",
    "    src/csv_reader.c src/csv_writer.c src/name_normalizer.c \\",
    "    src/country_normalizer.c src/rank_parser.c src/score_parser.c \\",
    "    src/requirement_parser.c src/utils.c -o build/clawer_normalizer.exe",
])
para("另針對 Windows 主控台預設使用 CP950（Big5）編碼導致 UTF-8 中文字串顯示為亂碼的問題，"
     "於 main.c 加入條件編譯處理：程式啟動時呼叫 SetConsoleOutputCP(CP_UTF8) 切換主控台編碼，"
     "並以 #ifdef _WIN32 包覆，確保 macOS / Linux 編譯不受影響。")

doc.add_heading("3.2 互動式操作流程", level=2)
para("主程式提供互動式 CLI 選單，並具備狀態防呆（未載入資料或未正規化時，擋下後續操作）：")
code_block([
    "==== Clawer C Data Normalization Engine ====",
    "1. 載入 CSV 資料並顯示原始資料",
    "2. 執行完整正規化流程並顯示結果",
    "3. 匯出正規化結果到 CSV 並離開程式",
    "0. 離開程式",
])
para("選單採同步複合動作設計：載入後立即顯示原始資料、正規化後立即顯示結果、"
     "匯出成功後自動離開程式。標準操作序列為 1 → 2 → 3，"
     "輸出檔為 data/samples/normalized_universities.csv。"
     "以下為 Windows 環境的實際執行畫面：")
IMG_DIR = r"C:\Users\User\OneDrive\Desktop\Clawer-C-Data-Normalization-Engine\c_engine\docs\img"
figure(IMG_DIR + r"\run-build-load.png",
       "圖 1：以 build.ps1 編譯（同步清空舊輸出）後執行，選項 1 載入 CSV 並顯示原始資料")
figure(IMG_DIR + r"\run-normalize-export.png",
       "圖 2：選項 2 正規化並顯示結果，選項 3 匯出 CSV 後自動離開程式")

doc.add_heading("3.3 正規化規則與實例", level=2)
make_table(
    ["模組", "原始輸入（raw）", "正規化輸出（normalized）"],
    [
        ("名稱", "University of   California Berkeley (UCB)", "University of California Berkeley (Ucb)"),
        ("名稱", "KAUST · hkust · eth", "KAUST · HKUST · ETH（縮寫保留全大寫）"),
        ("國家", "usa · U.S.A. · UK", "United States · United Kingdom"),
        ("國家", "The Netherlands · PR China", "Netherlands · China (Mainland)"),
        ("排名", "Top 100 · #10 · =201", "[1,100] · [10,10] · [201,201]"),
        ("排名", "101-150 · 150-101", "[101,150]（反轉區間自動交換）"),
        ("分數", '" 91.7 " · ".5"', "91.70 · 0.50（前導小數點修正）"),
        ("分數", '"" · "N/A" · "abc"', "-1（哨兵值，輸出層轉空欄）"),
    ],
    [1.6, 7.2, 7.2], mono_cols=(1, 2))
para("規則涵蓋 6 種排名格式與大小寫不敏感比對（\"TOP 100\" 等同 \"Top 100\"），"
     "並拒絕負數排名、TOP 0 與整數溢位（以 strtol 搭配合理上限防護）。")

doc.add_heading("3.4 弱點導向測試方法論", level=2)
para("本專題採弱點導向的系統性實驗測試，每輪 Session 依四個步驟進行："
     "（1）靜態分析——通讀模組程式碼並標記可疑邏輯；"
     "（2）探針程式——撰寫 probe 重現可疑行為；"
     "（3）修正——以最小修改修復根因；"
     "（4）回歸驗證——新增 regression 測試套件並全量重跑。"
     "累計完成六輪 Session，共 17 項修正（代號 A–T）：")
make_table(
    ["代號", "Session", "模組", "修正說明"],
    [
        ("A", "1/2", "country_normalizer", "新增 USA/UK/SG 等基本別名"),
        ("B", "1/2", "country_normalizer", "China/Hong Kong/Taiwan/Korea 別名"),
        ("C", "2", "rank_parser", "#N 井字號前綴格式支援"),
        ("D", "2", "rank_parser", "負數排名拒絕"),
        ("E", "3", "country_normalizer", "含連字號國名（Timor-Leste）保護"),
        ("F", "3", "utils", "title case 連接詞清單補全"),
        ("G", "3", "name_normalizer", "移除 Tech→Technology 破壞性展開"),
        ("H", "3", "csv_writer", "排名空欄輸出逗號數修正"),
        ("I", "4", "rank_parser", "TOP N / RANK N 大小寫不敏感"),
        ("J", "4", "country_normalizer", "Korea Rep. 等 19 個別名"),
        ("K", "4", "name_normalizer / utils", "ETH/UBC/HKUST/TUM 等縮寫擴充"),
        ("L", "4", "rank_parser", "sscanf 整數溢位 → strtol + 上限"),
        ("M", "5", "csv_reader", "rank 欄位部分空白智慧組合"),
        ("P", "5", "rank_parser", "TOP 0 拒絕（N ≥ 1）"),
        ("Q", "5", "rank_parser", "反轉區間 B-A 自動交換"),
        ("R", "6", "country_normalizer", "The X 前綴 + PR China + 38 個官方長名"),
        ("S", "6", "score_parser", "前導小數點 .5 → 0.5（原誤判 5.0）"),
        ("T", "6", "name_normalizer / utils", "KAUST / HKU 縮寫擴充"),
    ],
    [1.2, 2.0, 4.6, 8.2], mono_cols=(0,))

doc.add_heading("3.5 測試結果", level=2)
para("九個測試套件（make test_all）共 602 項測試全數通過：")
make_table(
    ["測試套件", "通過 / 總計", "驗證重點"],
    [
        ("test_normalizer", "17 / 17", "各模組基礎功能與 CSV 讀寫 round-trip"),
        ("test_extreme", "139 / 139", "極端與邊界輸入"),
        ("test_scale", "31 / 31", "資料規模與長字串"),
        ("test_weakness", "43 / 43", "弱點導向探針案例"),
        ("test_regression2–6", "372 / 372", "Sessions 2–6 修正之回歸防護"),
        ("合計", "602 / 602", "—"),
    ],
    [4.4, 3.2, 8.4], mono_cols=(0, 1))

# =====================================================================
# 4. Conclusion
# =====================================================================
doc.add_heading("4. Conclusion｜結論", level=1)

doc.add_heading("4.1 主要成果", level=2)
bullets([
    "完成 C11 正規化引擎：CSV Reader → 四大解析模組 → CSV Writer 的完整 pipeline，"
    "可將異質爬蟲輸出轉換為可直接比較的標準化紀錄。",
    "602 項單元／回歸測試全數通過，涵蓋基礎功能、極端值、資料規模與弱點情境。",
    "六輪弱點導向修正（A–T 共 17 項）建立了完整的回歸防護網，"
    "每項修正均有對應的 regression 測試確保不會再發。",
    "支援跨平台建置與執行（macOS / Linux 的 make 流程與 Windows 的 gcc 直接編譯），"
    "並解決 Windows 主控台中文編碼問題。",
])

doc.add_heading("4.2 已知限制", level=2)
bullets([
    "ASCII-only：名稱與國家正規化僅處理 ASCII 字元，É / ä / ã 等非 ASCII 字元直接傳遞，"
    "不做 Unicode 正規化（NFC/NFD）。",
    "requirement_parser 尚未實作，GMAT / GRE / GPA / IELTS / TOEFL / Duolingo "
    "等入學要求欄位目前不做解析。",
    "%.2f 捨入行為：score = 99.999 會輸出為 100.00（四捨五入超出滿分），"
    "屬 C 標準庫正常行為，暫以格式一致性為優先。",
])

doc.add_heading("4.3 未來工作", level=2)
bullets([
    "引入 utf8proc / ICU 等 Unicode 函式庫，處理非 ASCII 字元正規化。",
    "國家別名表外部化：改為自外部 mapping 檔載入，取代硬編碼，建立可維護的更新流程。",
    "與上游爬蟲系統整合並導入 CI 自動化流程（爬蟲 → 正規化引擎 → 下游資料庫），"
    "同時提供 library 形式介面（libnormalizer）。",
    "效能基準測試（吞吐量、記憶體佔用），對照 Python/pandas 實作驗證 C 設計選擇的優勢。",
])

# =====================================================================
# References
# =====================================================================
doc.add_heading("References｜參考文獻", level=1)
refs = [
    "QS World University Rankings, https://www.topuniversities.com/",
    "Times Higher Education (THE) World University Rankings, https://www.timeshighereducation.com/",
    "ShanghaiRanking / Academic Ranking of World Universities (ARWU), https://www.shanghairanking.com/",
    "ISO 3166-1, Codes for the representation of names of countries and their subdivisions.",
    "RFC 4180, Common Format and MIME Type for Comma-Separated Values (CSV) Files.",
    "Unicode Standard Annex #15, Unicode Normalization Forms.",
    "ISO/IEC 9899:2011, Information technology — Programming languages — C (C11).",
    "GCC, the GNU Compiler Collection — Documentation, https://gcc.gnu.org/onlinedocs/",
    "GNU Make Manual, https://www.gnu.org/software/make/manual/",
    "本專案內部文件：README.md、c_engine/docs/architecture.md、test-report.md（Session 1–6 累計測試報告）。",
]
for ref in refs:
    p = doc.add_paragraph(style='List Number')
    p.paragraph_format.space_after = Pt(3)
    set_run(p.add_run(ref), size=10.5)

OUT = r"C:\Users\User\OneDrive\Desktop\Clawer-C-Data-Normalization-Engine\b11402037-b11402053-report.docx"
doc.save(OUT)
print("saved:", OUT)
